import streamlit as st
import pandas as pd
from lifelines import KaplanMeierFitter, CoxPHFitter
from lifelines.statistics import logrank_test
import matplotlib.pyplot as plt
from docx import Document
from io import BytesIO

def generate_docx(median_survivals, hr, p_value, group_variable, survival_time_column):
    doc = Document()
    doc.add_heading('Kaplan-Meier Analysis Results', 0)
    
    for group, median_survival in median_survivals.items():
        doc.add_paragraph(f"Median {survival_time_column} for {group_variable}={group}: {median_survival:.2f} months")
    
    if p_value != "NA":
        doc.add_paragraph(f"P-value from Logrank Test: {p_value:.4f}")
    
    hr_str = ', '.join([f"{var}: {value:.4f}" for var, value in hr.items()])
    doc.add_paragraph(f"Hazard Ratio (HR) for {group_variable}: {hr_str}")
    
    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    
    return buffer.getvalue()

st.title('Kaplan-Meier Survival Analysis App')

uploaded_file = st.file_uploader("Choose a CSV file", type=['csv'])
if uploaded_file is not None:
    data = pd.read_csv(uploaded_file)
    column_names = data.columns.tolist()
    
    group_variable = st.selectbox("Select the group variable", column_names)
    event_column = st.selectbox("Select the event column", column_names)
    survival_time_column = st.selectbox("Select the survival time column", column_names)

    if st.button('Perform Kaplan-Meier Analysis'):
        kmf = KaplanMeierFitter()
        median_survivals = {}
        unique_groups = data[group_variable].unique()

        fig, ax = plt.subplots()
        for group in unique_groups:
            group_data = data[data[group_variable] == group]
            T = group_data[survival_time_column]
            E = group_data[event_column]

            kmf.fit(T, E, label=f'{group_variable}={group}')
            median_survivals[group] = kmf.median_survival_time_
            kmf.plot_survival_function(ax=ax, ci_show=False)

        plt.title(f'Kaplan-Meier Survival Analysis: Impact of {group_variable} on {survival_time_column}')
        plt.xlabel('Time')
        plt.ylabel('Survival Probability')
        plt.legend()
        st.pyplot(fig)

        # Cox Proportional Hazards model for Hazard Ratio
        cph = CoxPHFitter()
        cph_data = data[[group_variable, survival_time_column, event_column]].rename(columns={event_column: 'event'})
        cph.fit(cph_data, survival_time_column, event_col='event')
        hr = cph.summary['exp(coef)']

        # Logrank Test for p-value if there are exactly 2 groups
        p_value = "NA"
        if len(unique_groups) == 2:
            group1_data = data[data[group_variable] == unique_groups[0]]
            group2_data = data[data[group_variable] == unique_groups[1]]
            T1, E1 = group1_data[survival_time_column], group1_data[event_column]
            T2, E2 = group2_data[survival_time_column], group2_data[event_column]

            result = logrank_test(T1, T2, event_observed_A=E1, event_observed_B=E2)
            p_value = result.p_value

        # Generating and downloading DOCX report
        docx_bytes = generate_docx(median_survivals, hr, p_value, group_variable, survival_time_column)
        st.download_button(label="Download Analysis Report",
                           data=docx_bytes,
                           file_name="KM_Analysis_Report.docx",
                           mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document")
